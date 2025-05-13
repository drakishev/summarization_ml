import streamlit as st
import os
from langchain_ollama import ChatOllama, OllamaEmbeddings
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.prompts import ChatPromptTemplate
from llm import getStreamingChain, memory, remove_think_blocks
from models import get_list_of_models
import tempfile
from langchain_core.messages import HumanMessage, AIMessage
from docx import Document as DocxDocument  # Add this import
 
PERSIST_DIRECTORY = "storage"
EMBEDDING_MODEL = "nomic-embed-text"
SUMMARIZE_PROMPT = ChatPromptTemplate.from_template("""
### Instruction:
You are a helpful assistant that summarizes text clearly and concisely.
Provide a brief summary (8-10 sentences) of the provided text, capturing the main ideas.

## Text:    
{input_text}

## Summary:
""")

st.title("Local LLM with RAG and Document Summarization 📚")

# Initialize session state
if "list_of_models" not in st.session_state:
    st.session_state["list_of_models"] = get_list_of_models()
if "db" not in st.session_state:
    st.session_state.db = None
if "messages" not in st.session_state:
    st.session_state.messages = []
if "summary" not in st.session_state:
    st.session_state.summary = None

# Model selection
selected_model = st.sidebar.selectbox(
    "Select a model:",
    st.session_state["list_of_models"],
    index=st.session_state["list_of_models"].index("mistral") if "mistral" in st.session_state["list_of_models"] else 0
)
if st.session_state.get("ollama_model") != selected_model:
    st.session_state["ollama_model"] = selected_model
    st.session_state["llm"] = ChatOllama(
        model=selected_model,
        temperature=0.7,
        top_p=0.8,
        top_k=20,
        num_predict=512,
        model_kwargs={"enable_thinking": False}
    )

# Chunk size and overlap settings
chunk_size = st.sidebar.number_input("Chunk Size", min_value=200, max_value=2000, value=500, step=50)
chunk_overlap = st.sidebar.number_input("Chunk Overlap", min_value=0, max_value=500, value=50, step=10)

# File upload
st.subheader("Upload a Document")
uploaded_file = st.file_uploader("Choose a file (PDF, TXT, DOCX, MD)", type=["pdf", "txt", "docx", "md"])

@st.cache_resource
def load_and_process_file(uploaded_file):
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{uploaded_file.name.split('.')[-1]}") as tmp_file:
        tmp_file.write(uploaded_file.read())
        tmp_file_path = tmp_file.name

    if uploaded_file.name.endswith(".pdf"):
        loader = PyPDFLoader(tmp_file_path)
        docs = loader.load()
    elif uploaded_file.name.endswith(".docx"):
        # Use python-docx directly to load the document
        doc = DocxDocument(tmp_file_path)
        text = "\n".join([para.text for para in doc.paragraphs if para.text])
        docs = [Document(page_content=text, metadata={"source": uploaded_file.name, "page": "1"})]
    elif uploaded_file.name.endswith((".txt", ".md")):
        loader = TextLoader(tmp_file_path)
        docs = loader.load()
    else:
        st.error("Unsupported file type")
        return None

    os.unlink(tmp_file_path)
    return docs

@st.cache_resource
def summarize_text(text_list, _llm):
    full_text = "\n".join(text_list)
    full_text = full_text.replace("<think>", "").replace("</think>", "")
    if not full_text.strip():
        return "No text available for summarization."
    prompt = SUMMARIZE_PROMPT.format(input_text=full_text[:10000])
    response = _llm.invoke(prompt)
    summary = response.content if hasattr(response, "content") else response
    return remove_think_blocks(summary.replace("<think>", "").replace("</think>", ""))

@st.cache_resource
def load_documents_into_database(_docs, model_name, chunk_size, chunk_overlap):
    TEXT_SPLITTER = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    split_docs = TEXT_SPLITTER.split_documents(_docs)
    return Chroma.from_documents(
        documents=split_docs,
        embedding=OllamaEmbeddings(model=model_name),
        persist_directory=PERSIST_DIRECTORY,
        collection_name="rag_collection"
    )

if uploaded_file:
    docs = load_and_process_file(uploaded_file)
    if docs:
        text_list = [doc.page_content for doc in docs]
        st.session_state.summary = summarize_text(text_list, st.session_state["llm"])
        with st.spinner("Creating embeddings and loading into Chroma..."):
            st.session_state.db = load_documents_into_database(docs, EMBEDDING_MODEL, chunk_size, chunk_overlap)
        st.info("Document processed and ready for questions!")

if st.session_state.summary:
    st.subheader("Document Summary")
    st.markdown(st.session_state.summary)

@st.fragment
def chat_interface():
    st.subheader("Ask Questions About the Document")
    if st.session_state.db is None:
        st.warning("Please upload a document first.")
        st.chat_input("Question (document required)", disabled=True)
    else:
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        if prompt := st.chat_input("Question"):
            st.session_state.messages.append({"role": "user", "content": prompt})
            with st.chat_message("user"):
                st.markdown(prompt)
            with st.chat_message("assistant"):
                with st.spinner("Processing your question..."):
                    stream = getStreamingChain(
                        prompt,
                        st.session_state.messages,
                        st.session_state["llm"],
                        st.session_state.db,
                    )
                    response = ""
                    response_container = st.empty()
                    for chunk in stream:
                        cleaned_chunk = remove_think_blocks(chunk)
                        response += cleaned_chunk
                        response_container.markdown(response)
                    st.session_state.messages.append({"role": "assistant", "content": response})

chat_interface()
